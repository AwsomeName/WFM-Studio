#!/usr/bin/env python3
"""生成投标文件 .docx 模板，供 docx_write 工具引用。

模板包含：封面页、目录占位、各级标题样式、正文字体、
表格样式（表头蓝色底白字）、页眉页脚、预置章节骨架。

运行：python3 gen_template.py
输出：docs/templates/投标文件模板.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT = Path(__file__).resolve().parent.parent / "templates"
OUT.mkdir(parents=True, exist_ok=True)


def set_run_font(run, name_cn="宋体", name_en="Times New Roman", size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name_en
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), name_cn)
    if color:
        run.font.color.rgb = color


def add_paragraph(doc, text, name_cn="宋体", name_en="Times New Roman",
                  size=12, bold=False, alignment=None, color=None,
                  space_before=0, space_after=0):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, name_cn, name_en, size, bold, color)
    return p


def add_heading_styled(doc, text, level=1):
    """使用 Word 内置 Heading 样式，但覆盖字体。"""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    return h


def set_table_style(table):
    """给表格应用蓝色表头 + 边框样式。"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头行样式
    for cell in table.rows[0].cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2F5496"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    # 数据行样式
    for row_idx, row in enumerate(table.rows[1:], 1):
        bg = "F2F7FB" if row_idx % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            if row_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
                cell._tc.get_or_add_tcPr().append(shading)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_price_table(doc, title, headers, rows):
    """添加一个带标题的报价表格。"""
    add_heading_styled(doc, title, level=2)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    set_table_style(table)
    return table


def gen_bid_template():
    doc = Document()

    # ─── 全局默认样式 ──────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # ─── 页面设置 ──────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 页眉
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("投标文件 — {{项目名称}}")
    set_run_font(hr, "宋体", "Times New Roman", size=9,
                 color=RGBColor(0x80, 0x80, 0x80))

    # 页脚（页码）
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("第 ")
    set_run_font(fr, "宋体", "Times New Roman", size=9,
                 color=RGBColor(0x80, 0x80, 0x80))
    # 插入页码域代码
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    fp.runs[-1]._element.addnext(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar1.addnext(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    instrText.addnext(fldChar2)
    fr2 = fp.add_run(" 页 / 共 ")
    set_run_font(fr2, "宋体", "Times New Roman", size=9,
                 color=RGBColor(0x80, 0x80, 0x80))
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    fp.runs[-1]._element.addnext(fldChar3)
    instrText2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    fldChar3.addnext(instrText2)
    fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    instrText2.addnext(fldChar4)
    fr3 = fp.add_run(" 页")
    set_run_font(fr3, "宋体", "Times New Roman", size=9,
                 color=RGBColor(0x80, 0x80, 0x80))

    # ═══════════════════════════════════════════════════════════
    # 封面页
    # ═══════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    add_paragraph(doc, "投 标 文 件", name_cn="黑体", size=26, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  color=RGBColor(0x1F, 0x47, 0x88),
                  space_after=24)

    add_paragraph(doc, "（正本）", name_cn="宋体", size=16,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_after=36)

    # 封面信息表
    cover_fields = [
        ("项目名称", "{{项目名称}}"),
        ("项目编号", "{{项目编号}}"),
        ("招标单位", "{{招标单位}}"),
        ("投标单位", "{{投标单位}}"),
        ("投标日期", "{{投标日期}}"),
    ]
    cover_table = doc.add_table(rows=len(cover_fields), cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, placeholder) in enumerate(cover_fields):
        left = cover_table.rows[i].cells[0]
        right = cover_table.rows[i].cells[1]
        left.text = label
        right.text = placeholder
        for p in left.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                set_run_font(r, "黑体", "Times New Roman", size=14, bold=True)
        for p in right.paragraphs:
            for r in p.runs:
                set_run_font(r, "宋体", "Times New Roman", size=14,
                             color=RGBColor(0x80, 0x80, 0x80))

    # 移除封面表格边框
    for row in cover_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)

    # 分页
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 目录页（占位）
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "目  录", level=1)
    add_paragraph(doc, "（此处将由系统自动生成目录）", size=10,
                  color=RGBColor(0x80, 0x80, 0x80),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 第1章 投标函
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "第1章 投标函", level=1)

    add_paragraph(doc, "致：{{招标单位}}")
    add_paragraph(doc, "")
    add_paragraph(doc, "根据贵方发布的{{项目名称}}（项目编号：{{项目编号}}）招标文件，我方确认参与本次投标，并声明如下：")
    add_paragraph(doc, "")
    add_paragraph(doc, "一、我方已详细审阅招标文件全部内容，完全理解并接受其中的各项要求。")
    add_paragraph(doc, "二、我方投标总报价为人民币{{投标总报价（大写）}}（¥{{投标总报价（小写）}}）。")
    add_paragraph(doc, "三、我方承诺，若中标，将在合同签订后{{交付天数}}日内完成全部交付。")
    add_paragraph(doc, "四、本投标自开标之日起{{有效期天数}}天内有效。")
    add_paragraph(doc, "")
    add_paragraph(doc, "投标单位（盖章）：________________", space_before=24)
    add_paragraph(doc, "法定代表人（签字）：________________")
    add_paragraph(doc, "日期：________年____月____日")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 第2章 法定代表人授权书
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "第2章 法定代表人授权书", level=1)

    add_paragraph(doc, "本授权书声明：注册于{{公司注册地址}}的{{投标单位}}的{{法定代表人姓名}}代表本公司授权{{被授权人姓名}}为本公司的合法代理人，就{{项目名称}}的投标、谈判、签约及执行，以本公司名义处理一切与之有关的事务。")
    add_paragraph(doc, "")
    add_paragraph(doc, "本授权书自签字之日起生效，至上述事务全部完成之日终止。")
    add_paragraph(doc, "")
    add_paragraph(doc, "授权单位（盖章）：________________", space_before=24)
    add_paragraph(doc, "法定代表人（签字）：________________")
    add_paragraph(doc, "被授权人（签字）：________________")
    add_paragraph(doc, "日期：________年____月____日")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 第3章 报价明细表
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "第3章 报价明细表", level=1)

    price_headers = ["序号", "名称", "规格/型号", "单位", "数量", "单价(元)", "合价(元)"]

    add_price_table(doc, "表3-1 {{分类名称}}报价", price_headers, [
        ["1", "{{名称}}", "{{规格}}", "{{单位}}", "{{数量}}", "{{单价}}", "{{合价}}"],
        ["2", "...", "...", "...", "...", "...", "..."],
        ["", "", "", "", "小计", "", "{{小计}}"],
    ])

    add_paragraph(doc, "")
    add_paragraph(doc, "")

    add_price_table(doc, "表3-2 {{分类名称}}报价", price_headers, [
        ["1", "{{名称}}", "{{规格}}", "{{单位}}", "{{数量}}", "{{单价}}", "{{合价}}"],
        ["", "", "", "", "小计", "", "{{小计}}"],
    ])

    add_paragraph(doc, "")
    add_heading_styled(doc, "报价汇总", level=2)

    summary_table = doc.add_table(rows=4, cols=3)
    summary_table.rows[0].cells[0].text = "序号"
    summary_table.rows[0].cells[1].text = "报价项目"
    summary_table.rows[0].cells[2].text = "金额(元)"
    summary_table.rows[1].cells[0].text = "1"
    summary_table.rows[1].cells[1].text = "{{表3-1 分类名称}}"
    summary_table.rows[1].cells[2].text = "{{表3-1 小计}}"
    summary_table.rows[2].cells[0].text = "2"
    summary_table.rows[2].cells[1].text = "{{表3-2 分类名称}}"
    summary_table.rows[2].cells[2].text = "{{表3-2 小计}}"
    summary_table.rows[3].cells[0].text = ""
    summary_table.rows[3].cells[1].text = "报价总计"
    summary_table.rows[3].cells[2].text = "{{报价总计}}"
    set_table_style(summary_table)

    # 总计行加粗
    for cell in summary_table.rows[3].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 第4章 技术规格响应表
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "第4章 技术规格响应表", level=1)

    add_paragraph(doc, "以下逐项响应招标文件中的技术要求：")
    add_paragraph(doc, "")

    resp_table = doc.add_table(rows=4, cols=4)
    resp_table.rows[0].cells[0].text = "序号"
    resp_table.rows[0].cells[1].text = "招标文件要求"
    resp_table.rows[0].cells[2].text = "投标响应"
    resp_table.rows[0].cells[3].text = "偏离说明"
    resp_table.rows[1].cells[0].text = "1"
    resp_table.rows[1].cells[1].text = "{{招标技术要求1}}"
    resp_table.rows[1].cells[2].text = "{{投标响应内容1}}"
    resp_table.rows[1].cells[3].text = "{{无偏离/正偏离/负偏离}}"
    resp_table.rows[2].cells[0].text = "2"
    resp_table.rows[2].cells[1].text = "{{招标技术要求2}}"
    resp_table.rows[2].cells[2].text = "{{投标响应内容2}}"
    resp_table.rows[2].cells[3].text = "{{偏离说明}}"
    resp_table.rows[3].cells[0].text = "..."
    resp_table.rows[3].cells[1].text = "..."
    resp_table.rows[3].cells[2].text = "..."
    resp_table.rows[3].cells[3].text = "..."
    set_table_style(resp_table)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 第5章 资质证明文件
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "第5章 资质证明文件", level=1)
    add_paragraph(doc, "5.1 {{资质名称1}}（见附件1）")
    add_paragraph(doc, "5.2 {{资质名称2}}（见附件2）")
    add_paragraph(doc, "5.3 营业执照副本复印件（见附件3）")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 第6章 业绩证明
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "第6章 业绩证明", level=1)
    add_paragraph(doc, "近三年主要业绩：")
    add_paragraph(doc, "")

    perf_table = doc.add_table(rows=4, cols=5)
    perf_table.rows[0].cells[0].text = "序号"
    perf_table.rows[0].cells[1].text = "项目名称"
    perf_table.rows[0].cells[2].text = "业主单位"
    perf_table.rows[0].cells[3].text = "合同金额"
    perf_table.rows[0].cells[4].text = "完成年份"
    for r in range(1, 4):
        perf_table.rows[r].cells[0].text = str(r)
        for c in range(1, 5):
            perf_table.rows[r].cells[c].text = "{{...}}"
    set_table_style(perf_table)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 第7章 售后服务方案
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "第7章 售后服务方案", level=1)

    add_heading_styled(doc, "7.1 质保承诺", level=2)
    add_paragraph(doc, "{{质保期及质保范围}}")

    add_heading_styled(doc, "7.2 响应时间", level=2)
    add_paragraph(doc, "{{故障响应及到达现场时间承诺}}")

    add_heading_styled(doc, "7.3 备品备件", level=2)
    add_paragraph(doc, "{{备品备件供应方案}}")

    add_heading_styled(doc, "7.4 培训计划", level=2)
    add_paragraph(doc, "{{用户培训方案}}")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 第8章 项目实施计划
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "第8章 项目实施计划", level=1)

    add_heading_styled(doc, "8.1 生产进度", level=2)
    add_paragraph(doc, "{{生产排产计划及关键节点}}")

    add_heading_styled(doc, "8.2 质量控制", level=2)
    add_paragraph(doc, "{{质量控制体系及措施}}")

    add_heading_styled(doc, "8.3 交付方案", level=2)
    add_paragraph(doc, "{{交付方式、包装要求、运输方案}}")

    # ─── 保存 ─────────────────────────────────────────────────
    out_path = OUT / "投标文件模板.docx"
    doc.save(out_path)
    print(f"✅ 模板已生成: {out_path}")
    print(f"   大小: {out_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    gen_bid_template()
