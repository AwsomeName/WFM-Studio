#!/usr/bin/env python3
"""生成 WFM Studio 办公文档技能测试数据。

运行方式：python3 gen_test_data.py
输出目录：docs/samples/
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

OUT = Path(__file__).parent


# ─── helpers ──────────────────────────────────────────────────────────────

def set_cell_text(cell, text, bold=False, size=10):
    cell.text = str(text)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.name = "宋体"


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10)
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], val, size=10)
    return table


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "黑体"
    return h


def para(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(11)
        r.font.name = "宋体"
    return p


# ─── TC-1 / TC-6: 投标文件_v3.docx (含故意错误) ───────────────────────

def gen_bid_doc_v3():
    doc = Document()
    doc.core_properties.title = "投标文件 — 船用管路系统采购项目"
    doc.core_properties.author = "WFM Studio 测试"

    heading(doc, "投标文件", level=0)
    para(doc, "项目名称：船用管路系统采购项目")
    para(doc, "项目编号：ZB-2026-0342")
    para(doc, "投标单位：XX船舶配套有限公司")
    para(doc, "投标日期：2026年5月20日")

    heading(doc, "第1章 投标函", level=1)
    para(doc, "致：XX造船厂")
    para(doc, "根据贵方发布的船用管路系统采购项目（项目编号：ZB-2026-0342）招标文件，我方确认参与本次投标，并声明如下：")
    para(doc, "一、我方已详细审阅招标文件全部内容，完全理解并接受其中的各项要求。")
    para(doc, "二、我方投标总报价为人民币叁拾壹万玖仟零伍拾元整（¥319,050.00）。")
    para(doc, "三、我方承诺，若中标，将在合同签订后60日内完成全部交付。")
    para(doc, "四、本投标自开标之日起90天内有效。")

    heading(doc, "第2章 报价明细表", level=1)

    heading(doc, "表2-1 管材报价", level=2)
    add_table(doc,
        ["序号", "名称", "规格", "单位", "数量", "单价(元)", "合价(元)"],
        [
            ["1", "无缝钢管", "φ89×4.5", "吨", "12", "8,500", "102,000"],
            ["2", "不锈钢管", "φ60×3", "吨", "5", "28,000", "140,000"],
            ["3", "紫铜管", "φ22×1.5", "千克", "200", "85", "17,000"],
            ["", "", "", "", "小计", "", "259,000"],
        ])

    heading(doc, "表2-2 阀门报价", level=2)
    # ⚠️ 故意错误：止回阀 15×980=14,700 但写 14,000；小计应为 61,450 但写 60,050
    add_table(doc,
        ["序号", "名称", "规格", "单位", "数量", "单价(元)", "合价(元)"],
        [
            ["1", "蝶阀", "DN100-PN16", "个", "20", "1,200", "24,000"],
            ["2", "截止阀", "DN50-PN25", "个", "35", "650", "22,750"],
            ["3", "止回阀", "DN80-PN16", "个", "15", "980", "14,000"],       # ❌ 应为 14,700
            ["", "", "", "", "小计", "", "60,050"],                             # ❌ 应为 61,450
        ])

    para(doc, "")
    para(doc, "报价总计：319,050 元")  # ❌ 应为 320,450

    heading(doc, "第3章 技术规格响应", level=1)
    para(doc, "3.1 管材技术参数")
    para(doc, "无缝钢管符合 GB/T 8163 标准，材质 20# 钢，供货状态为热轧。")
    para(doc, "不锈钢管符合 GB/T 14976 标准，材质 316L，供货状态为固溶处理。")
    para(doc, "紫铜管符合 GB/T 1527 标准，材质 T2，供货状态为软态。")
    para(doc, "3.2 阀门技术参数")
    para(doc, "蝶阀：符合 GB/T 21385，阀体材质铸钢，阀板不锈钢，PN16。")
    para(doc, "截止阀：符合 GB/T 12233，阀体材质铸钢，PN25。")
    para(doc, "止回阀：符合 GB/T 12235，旋启式，阀体材质铸钢，PN16。")

    heading(doc, "第4章 资质证明文件", level=1)
    para(doc, "4.1 CCS 中国船级社工厂认可证书（有效期内）")
    para(doc, "4.2 ISO 9001 质量管理体系认证证书（有效期内）")
    para(doc, "4.3 营业执照副本复印件")

    heading(doc, "第5章 售后服务方案", level=1)
    para(doc, "5.1 质保期：验收合格后24个月。")
    para(doc, "5.2 响应时间：接到质量问题通知后，24小时内响应，48小时内到达现场。")
    para(doc, "5.3 备品备件：在质保期内免费提供损坏零部件的更换。")
    para(doc, "交付时间：合同签订后 60 天。")  # TC-6 选区测试目标文本

    doc.save(OUT / "投标文件_v3.docx")
    print("✅ 投标文件_v3.docx  (TC-1 金额核对 / TC-6 选区追问)")


# ─── TC-2: 招标需求说明书 ────────────────────────────────────────────────

def gen_tender_requirement():
    doc = Document()
    doc.core_properties.title = "招标需求说明书 — 船用管路系统采购项目"
    doc.core_properties.author = "WFM Studio 测试"

    heading(doc, "招标需求说明书", level=0)
    para(doc, "项目名称：船用管路系统采购项目")
    para(doc, "项目编号：ZB-2026-0342")
    para(doc, "招标单位：XX造船厂")
    para(doc, "发布日期：2026年5月10日")

    heading(doc, "一、项目概况", level=1)
    para(doc, "本项目为某型散货船管路系统采购，包含管材、阀门、法兰及附件。船舶总吨位约 35,000 DWT，航行区域为无限航区。管路系统涵盖机舱冷却水系统、压载水系统、消防水系统及舱底水系统。")

    heading(doc, "二、采购范围", level=1)
    para(doc, "1. 无缝钢管、不锈钢管、紫铜管等管材")
    para(doc, "2. 蝶阀、截止阀、止回阀、安全阀等阀门")
    para(doc, "3. 平焊法兰、对焊法兰等连接件")
    para(doc, "4. 管路支架、吊架等附件")

    heading(doc, "三、交付要求", level=1)
    para(doc, "交付时间：合同签订后 60 天")
    para(doc, "交付地点：XX造船厂码头仓库")
    para(doc, "质保期：验收合格后 24 个月")

    heading(doc, "四、资质要求", level=1)
    para(doc, "1. 具有 CCS（中国船级社）工厂认可")
    para(doc, "2. ISO 9001 质量管理体系认证")
    para(doc, "3. 近三年同类项目业绩不少于 3 个")

    heading(doc, "五、评标方法", level=1)
    para(doc, "综合评分法（商务 30% + 技术 50% + 价格 20%）")

    heading(doc, "六、投标文件要求", level=1)
    para(doc, "投标文件应包含以下内容：")
    para(doc, "- 投标函")
    para(doc, "- 法定代表人授权书")
    para(doc, "- 报价明细表")
    para(doc, "- 技术规格响应表")
    para(doc, "- 资质证明文件")
    para(doc, "- 业绩证明")
    para(doc, "- 售后服务方案")

    heading(doc, "七、技术规格要求", level=1)
    para(doc, "7.1 管材要求")
    para(doc, "碳钢管符合 GB/T 8163，不锈钢管符合 GB/T 14976，铜管符合 GB/T 1527。所有管材需提供材质证明书及船级社产品检验证书。")
    para(doc, "7.2 阀门要求")
    para(doc, "蝶阀符合 GB/T 21385，截止阀符合 GB/T 12233，止回阀符合 GB/T 12235。阀门需通过船级社型式认可。")
    para(doc, "7.3 连接件要求")
    para(doc, "法兰符合 GB/T 9119（平焊）及 GB/T 9115（对焊），螺栓螺母符合 GB/T 5782/GB/T 6170。")
    para(doc, "7.4 检验与试验")
    para(doc, "所有产品出厂前需进行水压试验，试验压力为公称压力的 1.5 倍。管材需逐根进行外观尺寸检查。")

    heading(doc, "八、投标截止时间", level=1)
    para(doc, "2026年6月15日 14:00（北京时间）")
    para(doc, "递交地点：XX造船厂采购部办公室")

    doc.save(OUT / "招标需求说明书_船用管路系统.docx")
    print("✅ 招标需求说明书_船用管路系统.docx  (TC-2 标书生成)")


# ─── TC-3: 工程量清单 + 单价参考表 ──────────────────────────────────────

def gen_budget_xlsx():
    # --- 工程量清单 ---
    wb1 = Workbook()
    ws1 = wb1.active
    ws1.title = "工程量清单"

    header_font = Font(name="微软雅黑", bold=True, size=11, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="2F5496")
    cell_font = Font(name="宋体", size=10)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    center = Alignment(horizontal="center", vertical="center")

    headers = ["分段号", "工序", "材料规格", "数量", "单位"]
    for c, h in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=c, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = thin_border

    data = [
        ["HD01", "钢板切割", "AH36 δ=12mm", 150, "平方米"],
        ["HD01", "钢板切割", "AH36 δ=16mm", 80, "平方米"],
        ["HD01", "型材加工", "L100×10", 200, "米"],
        ["HD02", "钢板切割", "AH36 δ=12mm", 120, "平方米"],
        ["HD02", "钢板切割", "AH36 δ=20mm", 60, "平方米"],
        ["HD03", "型材加工", "L120×12", 180, "米"],
        ["HD03", "拼装焊接", "CO₂焊", 500, "米"],
    ]
    for r_idx, row in enumerate(data, 2):
        for c_idx, val in enumerate(row, 1):
            cell = ws1.cell(row=r_idx, column=c_idx, value=val)
            cell.font = cell_font
            cell.alignment = center
            cell.border = thin_border

    for col in ["A", "B", "C", "D", "E"]:
        ws1.column_dimensions[col].width = 16

    wb1.save(OUT / "工程量清单_船体分段.xlsx")
    print("✅ 工程量清单_船体分段.xlsx  (TC-3)")

    # --- 单价参考表 ---
    wb2 = Workbook()
    ws2 = wb2.active
    ws2.title = "单价参考表"

    headers2 = ["工序", "材料规格", "综合单价(元)"]
    for c, h in enumerate(headers2, 1):
        cell = ws2.cell(row=1, column=c, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = thin_border

    data2 = [
        ["钢板切割", "AH36 δ=12mm", 280],
        ["钢板切割", "AH36 δ=16mm", 320],
        ["钢板切割", "AH36 δ=20mm", 380],
        ["型材加工", "L100×10", 45],
        ["型材加工", "L120×12", 55],
        ["拼装焊接", "CO₂焊", 120],
    ]
    for r_idx, row in enumerate(data2, 2):
        for c_idx, val in enumerate(row, 1):
            cell = ws2.cell(row=r_idx, column=c_idx, value=val)
            cell.font = cell_font
            cell.alignment = center
            cell.border = thin_border

    for col in ["A", "B", "C"]:
        ws2.column_dimensions[col].width = 20

    wb2.save(OUT / "单价参考表_2026.xlsx")
    print("✅ 单价参考表_2026.xlsx  (TC-3)")


# ─── TC-4: 技术规格书 ───────────────────────────────────────────────────

def gen_tech_spec():
    doc = Document()
    doc.core_properties.title = "技术规格书 — 某型散货船管路安装"
    doc.core_properties.author = "WFM Studio 测试"

    heading(doc, "技术规格书 — 某型散货船管路安装", level=0)
    para(doc, "文件编号：SPEC-2026-ME-001")
    para(doc, "版本：Rev.0")
    para(doc, "生效日期：2026年5月1日")

    heading(doc, "1. 适用范围", level=1)
    para(doc, "本规格适用于某型散货船机舱区域管路系统的预制与安装工程，包括但不限于：冷却水系统、压载水系统、消防水系统、舱底水系统、燃油系统及润滑油系统。")
    para(doc, "管路系统设计压力：1.6 MPa")
    para(doc, "管路系统设计温度：≤ 200°C")

    heading(doc, "2. 管材要求", level=1)
    para(doc, "2.1 碳钢管")
    para(doc, "符合 GB/T 8163 标准，材质为 20# 钢，供货状态为热轧。用于冷却水、压载水、消防水及舱底水系统。管壁厚度计算按 CCS《钢质海船入级规范》第3篇第4章执行。")
    para(doc, "2.2 不锈钢管")
    para(doc, "符合 GB/T 14976 标准，材质为 316L，供货状态为固溶处理。用于介质有腐蚀性要求的管路。")
    para(doc, "2.3 铜管")
    para(doc, "符合 GB/T 1527 标准，材质为 T2，供货状态为软态。用于仪表管路及小口径润滑油管路。")

    heading(doc, "3. 安装要求", level=1)
    para(doc, "3.1 管路布置应满足 CCS《钢质海船入级规范》相关要求。")
    para(doc, "3.2 法兰连接螺栓预紧力矩符合 CB/T 3766 要求，力矩值按下表执行：")

    add_table(doc,
        ["螺栓规格", "预紧力矩 (N·m)"],
        [
            ["M12", "45~55"],
            ["M16", "90~110"],
            ["M20", "170~210"],
            ["M24", "290~360"],
        ])

    para(doc, "3.3 管路支架间距不超过下表规定：")
    add_table(doc,
        ["管径 (DN)", "直管段最大间距 (m)", "弯管段最大间距 (m)"],
        [
            ["DN15~DN25", "1.5", "1.0"],
            ["DN32~DN50", "2.0", "1.5"],
            ["DN65~DN100", "3.0", "2.0"],
            ["DN125~DN200", "4.0", "3.0"],
            ["DN200 以上", "4.5", "3.5"],
        ])

    para(doc, "3.4 管路穿越水密舱壁时，必须使用经过船级社认可的水密贯通件。")

    heading(doc, "4. 焊接要求", level=1)
    para(doc, "4.1 碳钢管对接焊采用手工电弧焊（SMAW）或 CO₂气体保护焊（GMAW）。")
    para(doc, "4.2 不锈钢管对接焊采用氩弧焊（GTAW）。")
    para(doc, "4.3 焊接材料：碳钢焊条 E4303（J422），不锈钢焊条 A022。")
    para(doc, "4.4 焊前预热：壁厚 ≥ 20mm 时预热温度 100~150°C。")

    heading(doc, "5. 检验要求", level=1)
    para(doc, "5.1 焊缝 100% 外观检查，外观质量符合 CB/T 3558 要求。")
    para(doc, "5.2 对接焊缝射线检测（RT）或超声波检测（UT）抽检比例 ≥ 20%。")
    para(doc, "5.3 系统压力试验：试验压力为设计压力的 1.5 倍。")
    para(doc, "5.4 试验介质：水压试验采用洁净水，保压时间不少于 30 分钟。")

    heading(doc, "6. 表面处理与涂装", level=1)
    para(doc, "6.1 碳钢管外表面除锈等级 Sa2.5（近白级）。")
    para(doc, "6.2 底漆：环氧富锌底漆一道，干膜厚度 ≥ 40μm。")
    para(doc, "6.3 面漆：环氧面漆两道，干膜厚度 ≥ 160μm。")
    para(doc, "6.4 总干膜厚度 ≥ 200μm。")

    heading(doc, "7. 标识要求", level=1)
    para(doc, "7.1 管路安装完成后，按系统图进行标识。")
    para(doc, "7.2 标识内容：介质名称、流向箭头、管径、系统编号。")
    para(doc, "7.3 标识方式：色标 + 文字标签，符合 GB 7231 要求。")

    doc.save(OUT / "技术规格书_管路安装.docx")
    print("✅ 技术规格书_管路安装.docx  (TC-4)")


# ─── TC-7: 标书初稿 + 修订稿 ───────────────────────────────────────────

def gen_bid_doc_pair():
    # --- 初稿 ---
    doc1 = Document()
    doc1.core_properties.title = "标书 — 初稿"
    doc1.core_properties.author = "WFM Studio 测试"

    heading(doc1, "投标文件（初稿）", level=0)
    para(doc1, "项目：某型散货船管路安装工程")
    para(doc1, "投标单位：XX船舶工程有限公司")

    heading(doc1, "第1章 投标函", level=1)
    para(doc1, "致：XX造船厂")
    para(doc1, "我方确认参与某型散货船管路安装工程的投标。")

    heading(doc1, "第2章 法定代表人授权书", level=1)
    para(doc1, "本授权书声明：注册于（地址）的（公司名称）的（法人姓名）代表本公司授权（被授权人姓名）为本公司的合法代理人，就某型散货船管路安装工程的投标、谈判、签约及执行，以本公司名义处理一切与之有关的事务。")

    heading(doc1, "第3章 报价明细表", level=1)
    add_table(doc1,
        ["序号", "名称", "规格", "单位", "数量", "单价(元)", "合价(元)"],
        [
            ["1", "无缝钢管", "φ89×4.5", "吨", "12", "8,500", "102,000"],
            ["2", "不锈钢管", "φ60×3", "吨", "5", "28,000", "140,000"],
            ["3", "紫铜管", "φ22×1.5", "千克", "200", "85", "17,000"],
            ["", "", "", "", "小计", "", "259,000"],
        ])

    heading(doc1, "第4章 技术规格响应", level=1)
    para(doc1, "4.1 管材响应")
    para(doc1, "我方提供的无缝钢管符合 GB/T 8163 标准，材质 20# 钢。不锈钢管符合 GB/T 14976，材质 316L。铜管符合 GB/T 1527，材质 T2。")
    para(doc1, "4.2 焊接工艺")
    para(doc1, "碳钢管对接焊采用手工电弧焊（SMAW），焊条型号 E4303（J422）。焊前进行工艺评定试验，焊接工艺规程（WPS）经船级社审核批准。")  # 初稿中的描述
    para(doc1, "4.3 检验试验")
    para(doc1, "焊缝100%外观检查，对接焊缝RT/UT抽检比例≥20%。系统压力试验为设计压力的1.5倍，保压时间不少于30分钟。")

    heading(doc1, "第5章 资质证明文件", level=1)
    para(doc1, "5.1 CCS 工厂认可证书")
    para(doc1, "5.2 ISO 9001 质量管理体系认证")
    para(doc1, "5.3 营业执照副本")

    heading(doc1, "第6章 业绩证明", level=1)
    para(doc1, "近三年主要业绩：")
    para(doc1, "1. 35,000DWT散货船管路安装工程 — XX船厂（2024年）")
    para(doc1, "2. 50,000DWT油船管路安装工程 — YY船厂（2023年）")

    heading(doc1, "第7章 售后服务方案", level=1)
    para(doc1, "7.1 质保承诺：验收合格后24个月。")
    para(doc1, "7.2 响应时间：接到质量问题通知后24小时内响应，48小时内到达现场。")
    para(doc1, "7.3 备品备件：在质保期内免费提供损坏零部件的更换。")
    para(doc1, "7.4 培训计划：为船厂操作人员提供2天的管路系统操作维护培训。")

    heading(doc1, "第8章 项目实施计划", level=1)
    para(doc1, "8.1 生产进度：合同签订后第1周开始备料，第4周开始预制，第8周完成全部安装。")
    para(doc1, "8.2 质量控制：建立三级检验制度（自检、互检、专检）。")
    para(doc1, "8.3 交付方案：分批交付，每批附带检验报告和材质证明书。")

    doc1.save(OUT / "标书_初稿.docx")
    print("✅ 标书_初稿.docx  (TC-7)")

    # --- 修订稿 ---
    doc2 = Document()
    doc2.core_properties.title = "标书 — 修订稿"
    doc2.core_properties.author = "WFM Studio 测试"

    heading(doc2, "投标文件（修订稿）", level=0)
    para(doc2, "项目：某型散货船管路安装工程")
    para(doc2, "投标单位：XX船舶工程有限公司")
    para(doc2, "修订日期：2026年5月22日")

    heading(doc2, "第1章 投标函", level=1)
    para(doc2, "致：XX造船厂")
    para(doc2, "我方确认参与某型散货船管路安装工程的投标。")

    heading(doc2, "第2章 法定代表人授权书", level=1)
    para(doc2, "本授权书声明：注册于（地址）的（公司名称）的（法人姓名）代表本公司授权（被授权人姓名）为本公司的合法代理人，就某型散货船管路安装工程的投标、谈判、签约及执行，以本公司名义处理一切与之有关的事务。")

    heading(doc2, "第3章 报价明细表", level=1)
    # 修订：无缝钢管数量 12→15，新增弯头行，小计更新
    add_table(doc2,
        ["序号", "名称", "规格", "单位", "数量", "单价(元)", "合价(元)"],
        [
            ["1", "无缝钢管", "φ89×4.5", "吨", "15", "8,500", "127,500"],     # 12→15
            ["2", "不锈钢管", "φ60×3", "吨", "5", "28,000", "140,000"],
            ["3", "紫铜管", "φ22×1.5", "千克", "200", "85", "17,000"],
            ["4", "弯头 90°", "φ89 DN80", "个", "50", "45", "2,250"],           # 新增行
            ["", "", "", "", "小计", "", "286,750"],                               # 更新小计
        ])

    heading(doc2, "第4章 技术规格响应", level=1)
    para(doc2, "4.1 管材响应")
    para(doc2, "我方提供的无缝钢管符合 GB/T 8163 标准，材质 20# 钢。不锈钢管符合 GB/T 14976，材质 316L。铜管符合 GB/T 1527，材质 T2。")
    para(doc2, "4.2 焊接工艺")
    para(doc2, "碳钢管对接焊采用 CO₂气体保护焊（GMAW），焊丝型号 ER50-6，保护气体为 80%Ar+20%CO₂。焊前进行工艺评定试验，焊接工艺规程（WPS）经船级社审核批准。")  # 修改：手工电弧焊→CO₂焊
    para(doc2, "4.3 检验试验")
    para(doc2, "焊缝100%外观检查，对接焊缝RT/UT抽检比例≥20%。系统压力试验为设计压力的1.5倍，保压时间不少于30分钟。")

    heading(doc2, "第5章 资质证明文件", level=1)
    para(doc2, "5.1 CCS 工厂认可证书")
    para(doc2, "5.2 ISO 9001 质量管理体系认证")
    para(doc2, "5.3 营业执照副本")

    heading(doc2, "第6章 业绩证明", level=1)
    para(doc2, "近三年主要业绩：")
    para(doc2, "1. 35,000DWT散货船管路安装工程 — XX船厂（2024年）")
    para(doc2, "2. 50,000DWT油船管路安装工程 — YY船厂（2023年）")

    heading(doc2, "第7章 售后服务方案", level=1)
    para(doc2, "7.1 质保承诺：验收合格后24个月。")
    para(doc2, "7.2 响应时间：接到质量问题通知后24小时内响应，48小时内到达现场。")
    para(doc2, "7.3 备品备件：在质保期内免费提供损坏零部件的更换。")
    para(doc2, "7.4 培训计划：为船厂操作人员提供2天的管路系统操作维护培训。")
    # 新增 7.5
    para(doc2, "7.5 应急响应预案")
    para(doc2, "（一）应急组织架构：成立由项目经理任组长的应急响应小组，下设技术组、物资组、后勤组。")
    para(doc2, "（二）应急响应流程：接到应急通知后1小时内启动应急预案，2小时内技术人员出发，4小时内到达现场。")
    para(doc2, "（三）应急物资储备：常备密封垫片、阀门、法兰、管材等易损件，库存量不低于工程总量的5%。")

    heading(doc2, "第8章 项目实施计划", level=1)
    para(doc2, "8.1 生产进度：合同签订后第1周开始备料，第4周开始预制，第8周完成全部安装。")
    para(doc2, "8.2 质量控制：建立三级检验制度（自检、互检、专检）。")
    para(doc2, "8.3 交付方案：分批交付，每批附带检验报告和材质证明书。")

    doc2.save(OUT / "标书_修订稿.docx")
    print("✅ 标书_修订稿.docx  (TC-7)")


# ─── main ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("WFM Studio 测试数据生成\n" + "=" * 40)
    gen_bid_doc_v3()
    gen_tender_requirement()
    gen_budget_xlsx()
    gen_tech_spec()
    gen_bid_doc_pair()
    print("\n" + "=" * 40)
    print(f"全部完成！输出目录：{OUT}")
