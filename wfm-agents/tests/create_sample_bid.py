"""Create a sample bidding document for testing."""
from __future__ import annotations

from docx import Document

from pathlib import Path


def create_sample_bid(output_path: str | Path) -> Path:
    doc = Document()

    doc.add_heading("XX 船舶制造项目投标文件", level=1)
    doc.add_paragraph("项目编号：SB-2026-0042")
    doc.add_paragraph("投标单位：中船重工第 XXX 研究所")
    doc.add_paragraph("投标日期：2026 年 5 月")

    # --- Table 1: 工程量清单 ---
    doc.add_paragraph("表1 主要工程量清单")
    t1 = doc.add_table(rows=6, cols=5)
    t1.style = "Table Grid"
    headers = ["序号", "项目名称", "数量", "单价（元）", "合价（元）"]
    for ci, h in enumerate(headers):
        t1.rows[0].cells[ci].text = h

    rows = [
        ["1", "船体钢材", "120", "5,200.00", "624,000.00"],
        ["2", "焊接材料", "500", "380.00", "190,000.00"],
        ["3", "防腐涂料", "200", "1,500.00", "300,000.00"],
        ["4", "船用电缆", "1,000", "85.00", "85,000.00"],
        ["5", "管路附件", "300", "260.00", "78,000.00"],
    ]
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t1.rows[ri + 1].cells[ci].text = val

    # --- Table 2: 设备清单 (with an intentional error) ---
    doc.add_paragraph("表2 主要设备采购清单")
    t2 = doc.add_table(rows=5, cols=5)
    t2.style = "Table Grid"
    for ci, h in enumerate(headers):
        t2.rows[0].cells[ci].text = h

    rows2 = [
        ["1", "主推进电机", "2", "450,000.00", "900,000.00"],
        ["2", "导航雷达", "1", "280,000.00", "280,000.00"],
        ["3", "舵机系统", "4", "120,000.00", "480,000.00"],    # 错误：4×120000=480000，但下面写了 500000
        ["4", "通信系统", "2", "165,000.00", "330,000.00"],
    ]
    # Intentionally set wrong amount for row 3
    rows2[2] = ["3", "舵机系统", "4", "120,000.00", "500,000.00"]  # ❌ 480,000 ≠ 500,000
    for ri, row in enumerate(rows2):
        for ci, val in enumerate(row):
            t2.rows[ri + 1].cells[ci].text = val

    # --- Table 3: 汇总 ---
    doc.add_paragraph("表3 投标报价汇总表")
    t3 = doc.add_table(rows=5, cols=3)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "序号"
    t3.rows[0].cells[1].text = "费用项目"
    t3.rows[0].cells[2].text = "金额（元）"

    summary_rows = [
        ["1", "工程量清单合计", "1,277,000.00"],
        ["2", "设备采购合计", "2,010,000.00"],
        ["3", "措施项目费", "180,000.00"],
        ["4", "投标总价", "3,567,000.00"],    # ❌ 1277000+2010000+180000=3467000 ≠ 3567000
    ]
    for ri, row in enumerate(summary_rows):
        for ci, val in enumerate(row):
            t3.rows[ri + 1].cells[ci].text = val

    doc.add_paragraph("以上报价含增值税（13%）。")

    output_path = Path(output_path)
    doc.save(str(output_path))
    return output_path


if __name__ == "__main__":
    import sys
    path = sys.argv[1] if len(sys.argv) > 1 else "/tmp/投标文件_测试.docx"
    create_sample_bid(path)
    print(f"已创建测试文档: {path}")
    print()
    print("文档中埋了 2 个错误：")
    print("  1. 表2 第3行：4 × 120,000 = 480,000 ≠ 500,000（差 20,000）")
    print("  2. 表3 总价：1,277,000 + 2,010,000 + 180,000 = 3,467,000 ≠ 3,567,000（差 100,000）")
