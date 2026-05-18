"""Tests for docx review routing — ChatRequest + _extract_docx_review_extras."""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document
from fastapi import HTTPException

from wfm_agents.routes.chat import (
    ChatRequest,
    _extract_docx_review_extras,
)


def _make_docx_in_workspace(workspace: Path) -> Path:
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "名称"
    table.rows[0].cells[1].text = "数量"
    table.rows[0].cells[2].text = "金额"
    table.rows[1].cells[0].text = "钢材"
    table.rows[1].cells[1].text = "100"
    table.rows[1].cells[2].text = "500,000.00"
    target = workspace / "投标文件.docx"
    doc.save(str(target))
    return target


class TestExtractDocxReviewExtras:
    def test_valid_docx_path(self, tmp_path: Path) -> None:
        _make_docx_in_workspace(tmp_path)
        req = ChatRequest(
            workspace_root=str(tmp_path),
            message="核对金额",
            docx_path="投标文件.docx",
        )
        result = _extract_docx_review_extras(req, tmp_path)
        assert result is not None
        assert "docx_content" in result
        assert "docx_source" in result
        tables = result["docx_content"]["tables"]
        assert len(tables) == 1
        assert tables[0]["rows"][0][0] == "钢材"

    def test_no_docx_path_returns_none(self, tmp_path: Path) -> None:
        req = ChatRequest(
            workspace_root=str(tmp_path),
            message="你好",
        )
        result = _extract_docx_review_extras(req, tmp_path)
        assert result is None

    def test_file_not_found_raises_404(self, tmp_path: Path) -> None:
        req = ChatRequest(
            workspace_root=str(tmp_path),
            message="核对金额",
            docx_path="不存在的文件.docx",
        )
        with pytest.raises(HTTPException) as exc_info:
            _extract_docx_review_extras(req, tmp_path)
        assert exc_info.value.status_code == 404

    def test_non_docx_file_raises_400(self, tmp_path: Path) -> None:
        (tmp_path / "test.txt").write_text("hello")
        req = ChatRequest(
            workspace_root=str(tmp_path),
            message="核对金额",
            docx_path="test.txt",
        )
        with pytest.raises(HTTPException) as exc_info:
            _extract_docx_review_extras(req, tmp_path)
        assert exc_info.value.status_code == 400

    def test_path_traversal_raises_400(self, tmp_path: Path) -> None:
        req = ChatRequest(
            workspace_root=str(tmp_path),
            message="核对金额",
            docx_path="../../etc/passwd",
        )
        with pytest.raises(HTTPException) as exc_info:
            _extract_docx_review_extras(req, tmp_path)
        assert exc_info.value.status_code == 400

    def test_message_token_triggers_review(self, tmp_path: Path) -> None:
        _make_docx_in_workspace(tmp_path)
        req = ChatRequest(
            workspace_root=str(tmp_path),
            message="请核对 投标文件.docx 中的金额",
        )
        result = _extract_docx_review_extras(req, tmp_path)
        assert result is not None
        assert "docx_content" in result

    def test_message_token_quoted(self, tmp_path: Path) -> None:
        _make_docx_in_workspace(tmp_path)
        req = ChatRequest(
            workspace_root=str(tmp_path),
            message='核对 "投标文件.docx" 的金额',
        )
        result = _extract_docx_review_extras(req, tmp_path)
        assert result is not None

    def test_message_token_nonexistent_returns_none(self, tmp_path: Path) -> None:
        req = ChatRequest(
            workspace_root=str(tmp_path),
            message="核对 不存在的文件.docx 中的金额",
        )
        result = _extract_docx_review_extras(req, tmp_path)
        assert result is None

    def test_explicit_docx_path_takes_priority(self, tmp_path: Path) -> None:
        _make_docx_in_workspace(tmp_path)
        # Also create another docx
        doc = Document()
        doc.add_paragraph("另一个文件")
        (tmp_path / "other.docx")
        doc.save(str(tmp_path / "other.docx"))

        req = ChatRequest(
            workspace_root=str(tmp_path),
            message="请核对 other.docx 中的金额",
            docx_path="投标文件.docx",  # explicit should win
        )
        result = _extract_docx_review_extras(req, tmp_path)
        assert result is not None
        assert "投标文件.docx" in result["docx_source"]
