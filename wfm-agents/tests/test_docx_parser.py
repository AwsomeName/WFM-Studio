"""Tests for docx/parser.py — use python-docx to build fixture documents."""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches

from wfm_agents.docx.parser import (
    _expand_merged_cells,
    extract_amounts_from_table,
    format_docx_content,
    parse_docx,
)


def _make_docx(
    paragraphs: list[str] | None = None,
    tables: list[dict] | None = None,
) -> Path:
    """Create a temp .docx with given paragraphs and tables.

    tables format: [{"headers": [...], "rows": [[...], ...]}, ...]
    """
    doc = Document()
    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)
    if tables:
        for tbl in tables:
            headers = tbl.get("headers", [])
            rows = tbl.get("rows", [])
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for ci, h in enumerate(headers):
                table.rows[0].cells[ci].text = h
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    table.rows[ri + 1].cells[ci].text = val
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)  # noqa: SIM115
    doc.save(tmp.name)
    tmp.close()
    return Path(tmp.name)


class TestParseDocx:
    def test_empty_document(self) -> None:
        path = _make_docx()
        result = parse_docx(path)
        assert result["paragraphs"] == []
        assert result["tables"] == []
        assert result["stats"]["paragraphs_total"] == 0
        assert result["stats"]["tables_total"] == 0

    def test_paragraphs(self) -> None:
        path = _make_docx(paragraphs=["Hello", "World"])
        result = parse_docx(path)
        texts = [p["text"] for p in result["paragraphs"]]
        assert "Hello" in texts
        assert "World" in texts
        assert result["stats"]["paragraphs_kept"] == 2

    def test_table_extraction(self) -> None:
        path = _make_docx(
            paragraphs=["工程量清单"],
            tables=[{
                "headers": ["序号", "名称", "数量", "单价", "合价"],
                "rows": [
                    ["1", "钢材", "100", "5,000.00", "500,000.00"],
                    ["2", "混凝土", "200", "800.00", "160,000.00"],
                ],
            }],
        )
        result = parse_docx(path)
        assert len(result["tables"]) == 1
        tbl = result["tables"][0]
        assert tbl["row_count"] == 2
        assert tbl["col_count"] == 5
        assert tbl["headers"][0] == "序号"
        assert tbl["rows"][0][0] == "1"

    def test_multiple_tables(self) -> None:
        path = _make_docx(tables=[
            {"headers": ["A", "B"], "rows": [["1", "2"]]},
            {"headers": ["X", "Y"], "rows": [["3", "4"]]},
        ])
        result = parse_docx(path)
        assert len(result["tables"]) == 2

    def test_truncation(self) -> None:
        tables = [
            {"headers": ["C"], "rows": [[str(i)]]} for i in range(60)
        ]
        path = _make_docx(tables=tables)
        result = parse_docx(path, max_tables=5)
        assert result["stats"]["tables_total"] >= 50
        assert result["stats"]["tables_kept"] == 5

    def test_metadata(self) -> None:
        path = _make_docx(paragraphs=["test"])
        result = parse_docx(path)
        assert "title" in result["metadata"]
        assert "author" in result["metadata"]

    def test_file_size_limit(self) -> None:
        path = _make_docx()
        # Override size check by patching stat
        import unittest.mock
        with unittest.mock.patch.object(Path, "stat") as mock_stat:
            mock_stat.return_value.st_size = 20 * 1024 * 1024
            with pytest.raises(ValueError, match="文件过大"):
                parse_docx(path)

    def test_table_caption(self) -> None:
        doc = Document()
        doc.add_paragraph("表1 工程量清单")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "名称"
        table.rows[0].cells[1].text = "金额"
        table.rows[1].cells[0].text = "钢材"
        table.rows[1].cells[1].text = "5000"
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)  # noqa: SIM115
        doc.save(tmp.name)
        tmp.close()
        result = parse_docx(Path(tmp.name))
        tbl = result["tables"][0]
        assert tbl["caption"] is not None
        assert "工程量清单" in tbl["caption"]


class TestExtractAmounts:
    def test_basic_amounts(self) -> None:
        table = {
            "headers": ["名称", "金额"],
            "rows": [
                ["钢材", "5,000.00"],
                ["混凝土", "800.00"],
            ],
        }
        amounts = extract_amounts_from_table(table)
        assert len(amounts) == 2
        assert amounts[0]["value"] == 5000.0
        assert amounts[1]["value"] == 800.0

    def test_currency_symbols(self) -> None:
        table = {
            "headers": ["金额"],
            "rows": [["￥1,200"], ["¥3,000.50"], ["500元"]],
        }
        amounts = extract_amounts_from_table(table)
        values = [a["value"] for a in amounts]
        assert 1200.0 in values
        assert 3000.5 in values
        assert 500.0 in values

    def test_non_amount_text_skipped(self) -> None:
        table = {
            "headers": ["名称"],
            "rows": [["钢材"], ["混凝土"]],
        }
        amounts = extract_amounts_from_table(table)
        assert len(amounts) == 0

    def test_header_propagated(self) -> None:
        table = {
            "headers": ["单价", "合价"],
            "rows": [["100.00", "1,000.00"]],
        }
        amounts = extract_amounts_from_table(table)
        headers = {a["header"] for a in amounts}
        assert "单价" in headers
        assert "合价" in headers


class TestFormatDocxContent:
    def test_basic_formatting(self) -> None:
        content = {
            "metadata": {"title": "投标文件", "author": "", "created": ""},
            "paragraphs": [{"index": 0, "style": "Normal", "text": "摘要文本"}],
            "tables": [{
                "index": 0,
                "caption": "工程量清单",
                "headers": ["名称", "金额"],
                "rows": [["钢材", "5000"]],
                "row_count": 1,
                "col_count": 2,
            }],
            "stats": {"paragraphs_total": 1, "tables_total": 1,
                      "paragraphs_kept": 1, "tables_kept": 1},
        }
        md = format_docx_content(content)
        assert "# 投标文件" in md
        assert "摘要文本" in md
        assert "**工程量清单**" in md
        assert "| 名称 | 金额 |" in md
        assert "| 钢材 | 5000 |" in md

    def test_truncation_notice(self) -> None:
        content = {
            "metadata": {},
            "paragraphs": [],
            "tables": [],
            "stats": {"paragraphs_total": 0, "tables_total": 10,
                      "paragraphs_kept": 0, "tables_kept": 5},
        }
        md = format_docx_content(content)
        assert "截断提示" in md
