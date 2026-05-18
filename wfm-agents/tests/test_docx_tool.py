"""Tests for docx_read @function_tool — invoke via FunctionTool.on_invoke_tool."""

from __future__ import annotations

import json
import tempfile
from pathlib import Path
from unittest.mock import MagicMock

from docx import Document

from wfm_agents.agent_v2.tools import docx_read


def _make_docx(headers: list[str] | None = None, rows: list[list[str]] | None = None) -> Path:
    doc = Document()
    if headers:
        table = doc.add_table(rows=1 + (len(rows) if rows else 0), cols=len(headers))
        for ci, h in enumerate(headers):
            table.rows[0].cells[ci].text = h
        if rows:
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    table.rows[ri + 1].cells[ci].text = val
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)  # noqa: SIM115
    doc.save(tmp.name)
    tmp.close()
    return Path(tmp.name)


def _mock_ctx(workspace_root: str) -> MagicMock:
    ctx = MagicMock()
    ctx.context.workspace_root = workspace_root
    return ctx


async def _invoke(path: str, extract_tables_only: bool = False, workspace_root: str = "/tmp") -> str:
    """Invoke the FunctionTool via its on_invoke_tool method."""
    args = json.dumps({"path": path, "extract_tables_only": extract_tables_only})
    ctx = _mock_ctx(workspace_root)
    return await docx_read.on_invoke_tool(ctx, args)


class TestDocxReadTool:
    def test_reads_valid_docx(self, tmp_path: Path) -> None:
        path = _make_docx(
            headers=["名称", "金额"],
            rows=[["钢材", "5,000.00"]],
        )
        target = tmp_path / "test.docx"
        target.write_bytes(path.read_bytes())

        import asyncio
        result = asyncio.get_event_loop().run_until_complete(
            _invoke("test.docx", workspace_root=str(tmp_path))
        )
        assert not result.startswith("Error:")
        assert "钢材" in result
        assert "5,000.00" in result

    def test_file_not_found(self, tmp_path: Path) -> None:
        import asyncio
        result = asyncio.get_event_loop().run_until_complete(
            _invoke("nonexistent.docx", workspace_root=str(tmp_path))
        )
        assert result.startswith("Error:")
        assert "不存在" in result

    def test_non_docx_file(self, tmp_path: Path) -> None:
        txt = tmp_path / "test.txt"
        txt.write_text("hello")
        import asyncio
        result = asyncio.get_event_loop().run_until_complete(
            _invoke("test.txt", workspace_root=str(tmp_path))
        )
        assert result.startswith("Error:")
        assert ".docx" in result

    def test_path_traversal(self, tmp_path: Path) -> None:
        import asyncio
        result = asyncio.get_event_loop().run_until_complete(
            _invoke("../../etc/passwd", workspace_root=str(tmp_path))
        )
        assert result.startswith("Error:")

    def test_extract_tables_only(self, tmp_path: Path) -> None:
        doc = Document()
        doc.add_paragraph("这段文字应该被跳过")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "名称"
        table.rows[0].cells[1].text = "金额"
        table.rows[1].cells[0].text = "钢材"
        table.rows[1].cells[1].text = "5000"
        target = tmp_path / "test.docx"
        doc.save(str(target))

        import asyncio
        result = asyncio.get_event_loop().run_until_complete(
            _invoke("test.docx", extract_tables_only=True, workspace_root=str(tmp_path))
        )
        assert "这段文字应该被跳过" not in result
        assert "钢材" in result
